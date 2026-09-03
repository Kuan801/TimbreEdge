#!/usr/bin/env python3
# =============================================================================
#  report_docx.py  -  把 evaluate.py 的 JSON 變成一份 Word 報告
#
#  用法
#  -----
#    python3 evaluate.py PLAY.WAV 素材資料夾/ --start 60 --count 24 \
#            --json report.json --plot report.png
#    python3 report_docx.py report.json --plot report.png --out 音色比對報告.docx \
#            --source "Iowa MIS 鋼琴 mf，C4~B4 共 12 音" \
#            --method "Teensy 4.1 實機，按 w 錄成 PLAY.WAV"
#
#  --- 兩個設計決定 ---------------------------------------------------------
#
#  1) 吃 JSON，不去解析 evaluate.py 印出來的表格。
#     那張表是給人看的，欄寬和文字隨時會為了好讀而調整。拿它當資料介面，
#     改一次對齊就會靜靜地壞掉 —— 而且壞掉的樣子是「數字錯位」，不是報錯。
#
#  2) 判讀標準從 evaluate.py import，不在這裡重抄一份。
#     抄兩份的下場是改了門檻只改一邊，報告上的顏色與結論仍然照舊，
#     那種錯不會有任何跡象。
#
#  報告只放「量到的數字」與「判讀標準」，不代替使用者下結論 ——
#  結論該由看報告的人依標準自己判斷，那也是指導教授會問的第一件事。
# =============================================================================

import argparse
import datetime
import json
import os
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from evaluate import CRITERIA, METRIC_KEYS, grade   # noqa: E402

# 淺色底。Word 的網底色碼是 RRGGBB 純數字，沒有 # 前綴。
FILL = {"good": "E8F5E9", "ok": "FFF8E1", "bad": "FFEBEE", "na": "F5F5F5"}
MEAN_FILL = "CFD8DC"
HDR_FILL = "ECEFF1"

CJK_FONT = "Microsoft JhengHei"   # 微軟正黑體；沒有的話 Word 會自己找替代


def set_cell_bg(cell, hexcolor):
    """python-docx 沒有網底 API，只能自己塞 w:shd。"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")          # 一定要 clear；solid 會整格變黑
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def set_run_font(run, size=None, bold=None, color=None):
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = CJK_FONT
    # 中文字要另外指定 eastAsia，不然 Word 會用預設字型排中文
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)


def para(doc, text="", size=10, bold=False, italic=False, color=None,
         align=None, after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    r.font.italic = italic
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    return p


def put_cell(cell, text, *, bold=False, fill=None, align=WD_ALIGN_PARAGRAPH.RIGHT,
             size=8):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    if fill:
        set_cell_bg(cell, fill)


def fmt(key, v):
    if v is None:
        return "—"
    return CRITERIA[key][1](v)


def short_name(key):
    """表頭要短，不然 9 欄塞不進 A4。"""
    n = CRITERIA[key][0]
    return n.replace(" r", "").replace("誤差", "").replace("相關性", "")


# --------------------------------------------------------------------------
def build(data, args):
    doc = Document()

    # A4 直式，邊界縮一點才塞得下 10 欄的表
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Cm(1.8)
    sec.top_margin = sec.bottom_margin = Cm(2.0)

    st = doc.styles["Normal"]
    st.font.name = CJK_FONT
    st.font.size = Pt(10)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)

    notes = data["notes"]

    # ---------------------------------------------------------------- 標題 --
    h = doc.add_heading(args.title, level=0)
    for r in h.runs:
        r.font.name = CJK_FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    para(doc, "產生時間：" + datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
         size=9, italic=True, after=12)

    # ------------------------------------------------------------ 量測條件 --
    doc.add_heading("一、量測條件", level=1)
    para(doc, f"原始素材：{args.source}")
    para(doc, f"合成方式：{args.method}")
    para(doc, f"合成檔：{os.path.basename(data['scale_file'])}"
              f"（{data['scale_seconds']} 秒，{data['bpm']} BPM）")
    para(doc, f"比對音數：{len(notes)} 個"
              f"（合成檔共 {data['count']} 個音，從 MIDI {data['start_midi']} 開始）")
    para(doc, "比對方式：把合成的半音階依節拍切成單音，與同音高的原始素材逐一比較。"
              "所有跟時間有關的指標都先對齊起音點 —— 不對齊的話，切點的誤差會污染每一項數字。",
         after=12)

    # ------------------------------------------------------------ 逐音對照 --
    doc.add_heading("二、逐音對照", level=1)
    para(doc, "綠底＝很好，黃底＝可接受，紅底＝超出範圍。判讀標準見第三節。",
         size=9, after=6)

    t = doc.add_table(rows=1, cols=1 + len(METRIC_KEYS))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    put_cell(hdr[0], "音", bold=True, fill=HDR_FILL, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, k in enumerate(METRIC_KEYS):
        put_cell(hdr[1 + i], short_name(k), bold=True, fill=HDR_FILL,
                 align=WD_ALIGN_PARAGRAPH.CENTER)

    for n in notes:
        cells = t.add_row().cells
        put_cell(cells[0], n["note"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        for i, k in enumerate(METRIC_KEYS):
            put_cell(cells[1 + i], fmt(k, n.get(k)), fill=FILL[grade(k, n.get(k))])

    cells = t.add_row().cells
    put_cell(cells[0], "平均", bold=True, fill=MEAN_FILL, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, k in enumerate(METRIC_KEYS):
        put_cell(cells[1 + i], fmt(k, data["mean"].get(k)), bold=True, fill=MEAN_FILL)

    n_all_good = sum(1 for n in notes
                     if all(grade(k, n.get(k)) == "good" for k in METRIC_KEYS))
    n_any_bad = sum(1 for n in notes
                    if any(grade(k, n.get(k)) == "bad" for k in METRIC_KEYS))
    para(doc, "", after=4)
    para(doc, f"共 {len(notes)} 個音：{n_all_good} 個全部指標都落在「很好」，"
              f"{n_any_bad} 個至少有一項超出可接受範圍。", after=12)

    # -------------------------------------------------------- 判讀標準 -----
    doc.add_heading("三、指標與判讀標準", level=1)
    para(doc, "這些門檻不是這份報告訂的，是 evaluate.py 既有的標準，"
              "所有評測共用同一套，數字才能跨次比較。", size=9, after=6)

    ct = doc.add_table(rows=1, cols=3)
    ct.style = "Table Grid"
    h2 = ct.rows[0].cells
    put_cell(h2[0], "指標", bold=True, fill=HDR_FILL, align=WD_ALIGN_PARAGRAPH.LEFT, size=9)
    put_cell(h2[1], "本次平均", bold=True, fill=HDR_FILL, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    put_cell(h2[2], "判讀標準", bold=True, fill=HDR_FILL, align=WD_ALIGN_PARAGRAPH.LEFT, size=9)
    for k in METRIC_KEYS:
        v = data["mean"].get(k)
        c = ct.add_row().cells
        put_cell(c[0], CRITERIA[k][0], align=WD_ALIGN_PARAGRAPH.LEFT, size=9)
        put_cell(c[1], fmt(k, v), bold=True, fill=FILL[grade(k, v)],
                 align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
        put_cell(c[2], CRITERIA[k][4], align=WD_ALIGN_PARAGRAPH.LEFT, size=9)

    para(doc, "", after=4)
    para(doc, "噪聲量與噪聲位置兩欄專門補其他指標的盲點：噪聲層大約在 −70 dB，"
              "LSD、頻譜圖、質心都看不到它，但耳朵聽得出來"
              "（少了呼吸感、弓噪、槌擊聲）。", size=9, after=12)

    # ---------------------------------------------------------------- 圖 ---
    sec_no = 4
    if args.plot and os.path.exists(args.plot):
        doc.add_heading(f"四、比較圖", level=1)
        doc.add_picture(args.plot, width=Cm(17))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        sec_no = 5

    # ------------------------------------------------------------ 逐音明細 --
    cn = {4: "四", 5: "五"}[sec_no]
    doc.add_heading(f"{cn}、逐音明細", level=1)
    for n in notes:
        bad = [k for k in METRIC_KEYS if grade(k, n.get(k)) == "bad"]
        ok = [k for k in METRIC_KEYS if grade(k, n.get(k)) == "ok"]
        para(doc, f"{n['note']}　　參考素材：{n['ref']}", bold=True, after=2)
        para(doc, "　".join(f"{CRITERIA[k][0]} {fmt(k, n.get(k))}" for k in METRIC_KEYS),
             size=9, after=2)
        if not bad and not ok:
            para(doc, "全部指標都落在「很好」的區間。", size=9, color="2E7D32", after=8)
        else:
            if bad:
                para(doc, "超出可接受範圍：" + "、".join(CRITERIA[k][0] for k in bad),
                     size=9, color="C62828", after=2)
            if ok:
                para(doc, "在可接受但非最佳區間：" + "、".join(CRITERIA[k][0] for k in ok),
                     size=9, color="EF6C00", after=8)
            else:
                para(doc, "", after=6)

    return doc


def main():
    ap = argparse.ArgumentParser(description="把 evaluate.py 的 JSON 變成 Word 報告")
    ap.add_argument("json", help="evaluate.py --json 產生的檔案")
    ap.add_argument("--plot", default=None, help="evaluate.py --plot 產生的 PNG")
    ap.add_argument("--out", default="音色比對報告.docx")
    ap.add_argument("--title", default="TimbreClone 音色複製客觀評測報告")
    ap.add_argument("--source", default="（未註明）", help="原始素材是什麼")
    ap.add_argument("--method", default="（未註明）", help="合成是怎麼產生的")
    args = ap.parse_args()

    with open(args.json, encoding="utf-8") as f:
        data = json.load(f)
    if not data.get("notes"):
        print("JSON 裡沒有任何配對成功的音，先確認 evaluate.py 的 --start 對不對",
              file=sys.stderr)
        sys.exit(1)

    build(data, args).save(args.out)
    print(f"已輸出 {args.out}（{len(data['notes'])} 個音）")


if __name__ == "__main__":
    main()
